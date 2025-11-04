"""Document ingestion, splitting and indexing utilities for the RAG pipeline.

This module uses LangChain's document loaders and text splitter to create
consistent chunks, then exposes helpers to build a FAISS vector store using
Sentence-Transformers embeddings (all-MiniLM-L6-v2).

Functions:
 - load_and_split_documents(raw_dir) -> List[Document]
 - build_vectorstore(docs, index_path) -> FAISS vectorstore (saved to index_path)
 - process_documents(raw_dir, index_path) -> vectorstore

"""
import os
from typing import List

import pdfplumber
import docx
import chromadb
from chromadb.config import Settings


# Supported extensions
SUPPORTED = {".pdf", ".txt", ".docx"}


def _choose_loader(path: str):
    ext = os.path.splitext(path)[1].lower()
    return ext if ext in SUPPORTED else None


class SimpleDocument:
    def __init__(self, page_content: str, metadata: dict = None):
        self.page_content = page_content
        self.metadata = metadata or {}


def load_and_split_documents(raw_dir: str = "data/raw", chunk_size: int = 500, chunk_overlap: int = 50) -> List[SimpleDocument]:
    """Load supported files under ``raw_dir`` and split them into chunks.

    Returns a list of SimpleDocument objects with metadata.source set.
    """
    docs: List[SimpleDocument] = []

    def splitter(text: str):
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start += chunk_size - chunk_overlap
        return chunks

    if not os.path.exists(raw_dir):
        return []

    for fname in sorted(os.listdir(raw_dir)):
        path = os.path.join(raw_dir, fname)
        if not os.path.isfile(path):
            continue
        ext = _choose_loader(path)
        if ext is None:
            continue
        try:
            text = ""
            if ext == ".txt":
                with open(path, "r", encoding="utf-8") as f:
                    text = f.read()
            elif ext == ".pdf":
                try:
                    with pdfplumber.open(path) as pdf:
                        parts = [p.extract_text() or "" for p in pdf.pages]
                        text = "\n".join(parts)
                except Exception as e:
                    print(f"Warning: failed to extract PDF {path}: {e}")
                    continue
            elif ext == ".docx":
                try:
                    doc = docx.Document(path)
                    text = "\n".join(p.text for p in doc.paragraphs)
                except Exception as e:
                    print(f"Warning: failed to extract DOCX {path}: {e}")
                    continue

            # split text into chunks
            chunks = splitter(text)
            for i, c in enumerate(chunks):
                docs.append(SimpleDocument(c, metadata={"source": path, "chunk_id": f"{os.path.basename(path)}::{i}"}))
        except Exception as e:
            print(f"Warning: failed to load/split {path}: {e}")

    print(f"Loaded and split into {len(docs)} chunks from {raw_dir}")
    return docs


def build_vectorstore(docs: List[SimpleDocument], index_path: str = "data/index"):
    """Build a Chroma collection from LangChain Documents and persist it.

    Returns (client, collection).
    """
    os.makedirs(index_path, exist_ok=True)
    persist_dir = os.path.abspath(index_path)

    # Prepare documents, metadatas and stable ids based on chunk_id
    texts = [d.page_content for d in docs]
    metadatas = [d.metadata or {} for d in docs]
    ids = [d.metadata.get("chunk_id", f"doc-{i}") for i, d in enumerate(docs)]

    # embedding model (for storing embeddings in chroma)
    try:
        from sentence_transformers.SentenceTransformer import SentenceTransformer
        st_model = SentenceTransformer("all-MiniLM-L6-v2")
        embeddings = st_model.encode(texts, convert_to_numpy=True)
    except Exception:
        # Fallback: fastembed (no torch)
        from fastembed import TextEmbedding
        embedder = TextEmbedding(model_name="BAAI/bge-small-en-v1.5")
        # fastembed returns generator of lists
        embeddings = [vec for vec in embedder.embed(texts)]

    # Use new Chroma persistent client API
    client = chromadb.PersistentClient(path=persist_dir)
    # create or get collection
    try:
        # delete existing collection to replace
        client.delete_collection(name="rag_collection")
    except Exception:
        pass
    collection = client.get_or_create_collection(name="rag_collection")

    collection.add(ids=ids, documents=texts, metadatas=metadatas, embeddings=embeddings.tolist())
    print(f"Saved Chroma collection to {persist_dir}")
    return client, collection


def process_documents(raw_dir: str = "data/raw", index_path: str = "data/index"):
    """High-level helper to load, split, embed and index documents.

    Returns (client, collection) on success or None if no docs found.
    """
    docs = load_and_split_documents(raw_dir)
    if not docs:
        print("No documents found to process.")
        return None
    client, collection = build_vectorstore(docs, index_path=index_path)
    return client, collection


def update_vectorstore_incremental(raw_dir: str = "data/raw", index_path: str = "data/index") -> int:
    """Incrementally update the Chroma collection with only new chunks.

    Returns the number of chunks added.
    """
    docs = load_and_split_documents(raw_dir)
    if not docs:
        return 0

    persist_dir = os.path.abspath(index_path)
    client = chromadb.PersistentClient(path=persist_dir)
    collection = client.get_or_create_collection(name="rag_collection")

    # embed model for new chunks
    try:
        from sentence_transformers.SentenceTransformer import SentenceTransformer
        st_model = SentenceTransformer("all-MiniLM-L6-v2")
        def encode_batch(tt: list):
            return st_model.encode(tt, convert_to_numpy=True).tolist()
    except Exception:
        from fastembed import TextEmbedding
        _fe = TextEmbedding(model_name="BAAI/bge-small-en-v1.5")
        def encode_batch(tt: list):
            return [v for v in _fe.embed(tt)]

    added = 0
    # Add in small batches to reduce round-trips
    batch_texts = []
    batch_metas = []
    batch_ids = []

    def flush_batch():
        nonlocal added, batch_texts, batch_metas, batch_ids
        if not batch_ids:
            return
        embeddings = encode_batch(batch_texts)
        collection.add(ids=batch_ids, documents=batch_texts, metadatas=batch_metas, embeddings=embeddings)
        added += len(batch_ids)
        batch_texts, batch_metas, batch_ids = [], [], []

    for i, d in enumerate(docs):
        cid = d.metadata.get("chunk_id", f"doc-{i}")
        try:
            # check existence by id
            got = collection.get(ids=[cid])
            if got and got.get("ids"):
                continue  # already present
        except Exception:
            # if get fails for any reason, attempt to add; duplicates will be rejected by backend
            pass
        batch_ids.append(cid)
        batch_texts.append(d.page_content)
        batch_metas.append(d.metadata or {})
        if len(batch_ids) >= 64:
            flush_batch()

    flush_batch()
    if added:
        print(f"Incrementally added {added} chunks to {persist_dir}")
    return added

